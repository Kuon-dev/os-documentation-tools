import os
from pygments.lexers import get_lexer_by_name, JavascriptLexer
from io import BytesIO
import sys
from typing import List, Tuple
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from pygments.util import ClassNotFound
from langchain_anthropic import ChatAnthropic
from langchain_core.output_parsers import StrOutputParser
from rich.console import Console
from rich.logging import RichHandler
from rich.table import Table
import logging
from dotenv import load_dotenv
from pathspec import PathSpec
from pathspec.patterns.gitwildmatch import GitWildMatchPattern
import os
from typing import List, Tuple, Optional
from pygments import highlight
from pygments.lexers import get_lexer_by_name
from pygments.formatters import ImageFormatter
from pygments.styles import get_style_by_name
from PIL import Image, ImageDraw, ImageFont
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn

from PIL import Image

# Load environment variables
load_dotenv()

# Set up rich logging
logging.basicConfig(
    level="INFO",
    format="%(message)s",
    datefmt="[%X]",
    handlers=[RichHandler(rich_tracebacks=True)]
)
log = logging.getLogger("rich")
console = Console()

# Constants for pricing
PRICE_PER_1M_INPUT_TOKENS = 0.25  # $3 per 1M input tokens
PRICE_PER_1M_OUTPUT_TOKENS = 1.25  # $15 per 1M output tokens

def get_gitignore_spec(directory: str) -> PathSpec:
    gitignore_path = os.path.join(directory, '.gitignore')
    if os.path.exists(gitignore_path):
        with open(gitignore_path, 'r') as gitignore_file:
            gitignore_content = gitignore_file.read()
        return PathSpec.from_lines(GitWildMatchPattern, gitignore_content.splitlines())
    return PathSpec([])

def read_typescript_files(directory: str) -> List[Tuple[str, str]]:
    gitignore_spec = get_gitignore_spec(directory)
    typescript_files = []
    
    for root, dirs, files in os.walk(directory):
        # Skip .git and other unnecessary directories
        dirs[:] = [d for d in dirs if d not in ['.git', 'node_modules', 'dist', 'build']]
        
        rel_path = os.path.relpath(root, directory)
        
        if "tests" in rel_path.split(os.path.sep) or "seeders" in rel_path.split(os.path.sep) or "config" in rel_path.split(os.path.sep):
            continue
        
        if gitignore_spec.match_file(rel_path):
            continue
        
        for file in files:
            if file.endswith('.tsx') or file.endswith('.ts'):
                file_path = os.path.join(root, file)
                rel_file_path = os.path.relpath(file_path, directory)
                if not gitignore_spec.match_file(rel_file_path):
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                            typescript_files.append((rel_file_path, content))
                    except Exception as e:
                        log.error(f"[red]Error reading file {file_path}: {str(e)}[/red]")
    
    log.info(f"[green]Read {len(typescript_files)} TypeScript files[/green]")
    return typescript_files

def create_output_directory():
    output_dir = "./output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    return output_dir

def generate_explanation(file_name: str, file_content: str) -> Tuple[str, str, int, int]:
    anthropic_api_key = os.getenv("ANTHROPIC_API_KEY")
    openai_api_key = os.getenv("OPENAI_API_KEY")
    
    llm = ChatOpenAI(model="gpt-4o-mini", openai_api_key=openai_api_key, max_tokens=4096)
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert in TypeScript and software development. Your task is to provide a brief explanation of the given TypeScript code and a short caption."),
        ("human", """
        Analyze the following TypeScript code and provide:
        1. A brief explanation (40-60 words) of its purpose and functionality
        2. A short caption (10 words or less) summarizing the code's main function

        File Name: {file_name}

        Code Content:
        {file_content}

        Format your response as follows:
        Caption: [Your 10-word or less caption here]

        [Your 40-60 word explanation here]

        Ensure the explanation is informative yet brief, suitable for a quick overview of the file's contents.
        """)
    ])
    
    chain = prompt | llm | StrOutputParser()
    
    try:
        result = chain.invoke({"file_name": file_name, "file_content": file_content})
        log.info(f"[green]Explanation and caption generated for {file_name}[/green]")
        
        # Split the result into caption and explanation
        caption, explanation = result.split("\n\n", 1)
        caption = caption.replace("Caption: ", "").strip()
        
        # Calculate token usage
        input_tokens = llm.get_num_tokens(prompt.format(file_name=file_name, file_content=file_content))
        output_tokens = llm.get_num_tokens(result)
        
        return caption, explanation, input_tokens, output_tokens
    except Exception as e:
        log.error(f"[red]Error generating explanation for {file_name}: {str(e)}[/red]")
        log.error(f"File content (first 100 characters): {file_content[:100]}")
        return "", "", 0, 0

def get_default_font():
    try:
        # Try to use a common system font
        return ImageFont.truetype("C:\\Users\\aaron\\Downloads\\CascadiaMono\\CaskaydiaMonoNerdFont-Regular.ttf", 16)
    except IOError:
        try:
            # Fallback to default
            return ImageFont.load_default()
        except IOError:
            # If all else fails, use a bitmap font
            return ImageFont.load_default()

def get_lexer_for_file(file_name: str):
    extension = os.path.splitext(file_name)[1][1:]  # Remove the dot
    if extension in ['tsx', 'jsx']:
        return JavascriptLexer()
    try:
        return get_lexer_by_name(extension)
    except ClassNotFound:
        return JavascriptLexer()  # Fallback to JavaScript lexer

def create_code_screenshot(file_name: str, code_content: str, caption: str, output_dir: str) -> str:
    try:
        # Determine the language based on file extension
        lexer = get_lexer_for_file(file_name)
        
        # Use a dark style for better readability
        style = get_style_by_name("dracula")
        
        # Ensure code_content is a string
        if isinstance(code_content, bytes):
            code_content = code_content.decode('utf-8', errors='replace')
        
        # Generate the syntax-highlighted image
        formatter = ImageFormatter(
            style=style,
            line_numbers=True,
            font_size=14,
            font_name="Cascadia Mono",
            line_number_separator=True,
            line_number_pad=3,
            hl_lines=[],
            line_number_bg="#282a36",  # Dracula background color
            line_number_fg="#6272a4"   # Dracula comment color for line numbers
        )

        highlighted_code = highlight(code_content, lexer, formatter)
        
        # Open the image using Pillow
        code_image = Image.open(BytesIO(highlighted_code))
        
        # Create a new image with extra space for the caption
        caption_height = 40
        new_image = Image.new('RGB', (code_image.width, code_image.height + caption_height), color='#272822')
        
        # Paste the code image onto the new image
        new_image.paste(code_image, (0, caption_height))
        
        # Add the caption
        draw = ImageDraw.Draw(new_image)
        font = get_default_font()
        draw.text((10, 10), f"{file_name}: {caption}", font=font, fill='#f8f8f2')
        
        # Save the image in the output directory
        screenshot_dir = os.path.join(output_dir, "screenshots")
        os.makedirs(screenshot_dir, exist_ok=True)
        
        # Use os.path.basename to get just the filename without the path
        base_name = os.path.basename(file_name)
        output_path = os.path.join(screenshot_dir, f"{os.path.splitext(base_name)[0]}.png")
        new_image.save(output_path)
        
        return os.path.relpath(output_path, output_dir)
    except Exception as e:
        log.error(f"[red]Error creating screenshot for {file_name}: {str(e)}[/red]")
        log.error(f"Code content type: {type(code_content)}")
        log.error(f"Code content (first 100 characters): {str(code_content[:100])}")
        return ""

def calculate_cost(input_tokens: int, output_tokens: int) -> float:
    input_cost = (input_tokens / 1_000_000) * PRICE_PER_1M_INPUT_TOKENS
    output_cost = (output_tokens / 1_000_000) * PRICE_PER_1M_OUTPUT_TOKENS
    total_cost = input_cost + output_cost
    return total_cost

def generate_explanation_with_screenshot(file_name: str, file_content: str, output_dir: str) -> Tuple[str, str, str, int, int]:
    try:
        caption, explanation, input_tokens, output_tokens = generate_explanation(file_name, file_content)
        screenshot_path = create_code_screenshot(file_name, file_content, caption, output_dir)
        return caption, explanation, screenshot_path, input_tokens, output_tokens
    except Exception as e:
        log.error(f"[red]Error in generate_explanation_with_screenshot for {file_name}: {str(e)}[/red]")
        return "", "", "", 0, 0

def resize_image(image_path, max_height_inches=9):
    with Image.open(image_path) as img:
        # Get the original width and height
        width, height = img.size
        
        # Calculate the aspect ratio
        aspect_ratio = width / height
        
        # Calculate new dimensions
        new_height = min(height, int(max_height_inches * 88))  # 96 DPI
        new_width = int(new_height * aspect_ratio)
        
        return new_width, new_height

def add_seq_field(run, seq_identifier):
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' SEQ {seq_identifier} \\* ARABIC'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def create_word_document(explanations: List[Tuple[str, str, str, str]], output_dir: str):
    try:
        doc = Document()
        doc.add_heading('Code Explanations', 0)

        # Create a new style for figure captions
        styles = doc.styles
        figure_style = styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
        figure_style.base_style = styles['Caption']
        figure_style.font.italic = True
        figure_style.font.size = Pt(10)

        for file_name, caption, explanation, screenshot_path in explanations:
            # Add file name as heading
            heading = doc.add_paragraph(f"File: {file_name}")
            heading.style = 'Heading 1'

            # Check if the screenshot file exists before adding it
            full_screenshot_path = os.path.join(output_dir, screenshot_path)
            if os.path.exists(full_screenshot_path):
                # Resize the image
                new_width, new_height = resize_image(full_screenshot_path)
                
                # Add a new paragraph for the image and center-align it
                image_paragraph = doc.add_paragraph()
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_run = image_paragraph.add_run()
                image_run.add_picture(full_screenshot_path, width=Inches(new_width/96), height=Inches(new_height/96))

                # Add figure caption with auto-numbering
                figure_caption = doc.add_paragraph()
                figure_caption.style = figure_style
                figure_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                figure_caption.add_run('Figure ')
                add_seq_field(figure_caption.add_run(), 'Figure')
                figure_caption.add_run(f': {caption}')
            else:
                log.warning(f"[yellow]Screenshot not found for {file_name}. Skipping image.[/yellow]")

            # Add justified explanation
            explanation_paragraph = doc.add_paragraph(explanation)
            explanation_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Add line spaces, except for the last item
            if explanations.index((file_name, caption, explanation, screenshot_path)) < len(explanations) - 1:
                for _ in range(2):
                    spacer = doc.add_paragraph()
                    spacer.space_after = Pt(12)  # Adjust this value to increase/decrease space

        # Save the document
        output_path = os.path.join(output_dir, "code_explanations.docx")
        doc.save(output_path)
        log.info(f"[green]Word document created with auto-numbered figures: {output_path}[/green]")
        return output_path
    except Exception as e:
        log.error(f"[red]Error creating Word document: {str(e)}[/red]")
        return None

def generate_code_explanations(directory: str, output_dir: str) -> Tuple[List[Tuple[str, str, str, str]], int, int, float, Optional[str]]:
    try:
        typescript_files = read_typescript_files(directory)
        explanations = []
        total_input_tokens = 0
        total_output_tokens = 0
        
        for file_name, file_content in typescript_files:
            try:
                caption, explanation, screenshot_path, input_tokens, output_tokens = generate_explanation_with_screenshot(file_name, file_content, output_dir)
                if explanation:
                    explanations.append((file_name, caption, explanation, screenshot_path))
                    total_input_tokens += input_tokens
                    total_output_tokens += output_tokens
            except Exception as e:
                log.error(f"[red]Error processing file {file_name}: {str(e)}[/red]")
        
        cost = calculate_cost(total_input_tokens, total_output_tokens)
        
        # Create Word document instead of saving to Markdown
        word_doc_path = create_word_document(explanations, output_dir)
        
        return explanations, total_input_tokens, total_output_tokens, cost, word_doc_path
    except Exception as e:
        log.error(f"[red]Error in main pipeline: {str(e)}[/red]")
        return [], 0, 0, 0.0, None

# Update the main execution
if __name__ == "__main__":
    console.print("[bold green]Code Explanation Generator for TypeScript[/bold green]")
    
    project_directory = os.getenv("PROJECT_DIRECTORY")
    if not project_directory:
        console.print("[bold red]PROJECT_DIRECTORY not found in environment variables[/bold red]")
        console.print("Please set the PROJECT_DIRECTORY in your .env file")
        sys.exit(1)
    
    console.print(f"[cyan]Analyzing project directory: {project_directory}[/cyan]")
    
    output_dir = create_output_directory()
    
    explanations, input_tokens, output_tokens, cost, word_doc_path = generate_code_explanations(project_directory, output_dir)
    if explanations:
        if word_doc_path:
            console.print("[bold green]Code explanations and Word document generated successfully![/bold green]")
            console.print(f"[bold green]Word document saved in: {word_doc_path}[/bold green]")
        else:
            console.print("[bold yellow]Code explanations generated, but Word document creation failed. Check logs for details.[/bold yellow]")
        
        # Display token usage and cost
        table = Table(title="Token Usage and Cost")
        table.add_column("Metric", style="cyan")
        table.add_column("Value", style="magenta")
        table.add_row("Input Tokens", str(input_tokens))
        table.add_row("Output Tokens", str(output_tokens))
        table.add_row("Total Tokens", str(input_tokens + output_tokens))
        table.add_row("Estimated Cost", f"${cost:.4f}")
        console.print(table)
    else:
        console.print("[bold red]Failed to generate code explanations. Check logs for details.[/bold red]")
